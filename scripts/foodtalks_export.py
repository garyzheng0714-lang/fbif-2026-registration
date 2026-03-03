#!/usr/bin/env python3
"""
Export FoodTalks article list to local DOCX files + Excel metadata table.

Features:
- Fetches paginated article list from FoodTalks API.
- Fetches per-article details (including HTML body).
- Writes one DOCX per article.
- Writes JSONL/CSV/XLSX metadata with source/author/filter fields.
- Supports resume by reusing existing metadata + DOCX files.
"""

from __future__ import annotations

import argparse
import json
import random
import re
import time
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

import pandas as pd
import requests
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document


API_BASE = "https://api-we.foodtalks.cn"
LIST_API = f"{API_BASE}/news/news/page"
DETAIL_API = f"{API_BASE}/news/news"
SITE_BASE = "https://www.foodtalks.cn"


DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/122.0.0.0 Safari/537.36"
    ),
    "Accept": "application/json, text/plain, */*",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Origin": SITE_BASE,
    "Referer": f"{SITE_BASE}/news",
}


BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote", "pre"}
CONTAINER_TAGS = {"body", "article", "section", "main", "div"}
SKIP_TAGS = {"script", "style", "noscript", "iframe"}
ILLEGAL_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def clean_excel_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    text = ILLEGAL_CONTROL_CHARS.sub("", text)
    if len(text) > 32767:
        text = text[:32767]
    return text


def clean_docx_text(value: Any) -> str:
    if value is None:
        return ""
    return ILLEGAL_CONTROL_CHARS.sub("", str(value))


def safe_filename(name: str, max_len: int = 120) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", name or "")
    cleaned = normalize_text(cleaned).strip(" .")
    if not cleaned:
        cleaned = "untitled"
    if len(cleaned) > max_len:
        cleaned = cleaned[:max_len].rstrip(" .")
    return cleaned


def as_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def is_fbif_original(source_name: str, source_type: Any, source_id: Any) -> bool:
    source = (source_name or "").lower()
    if "fbif" in source or "foodtalks" in source:
        return True
    if as_int(source_id) in {7630, 7632}:
        return True
    return as_int(source_type, -1) == 0 and source in {"foodtalks", "fbif食品饮料创新"}


def source_label(source_name: str, source_type: Any, source_id: Any) -> str:
    if is_fbif_original(source_name, source_type, source_id):
        return "FBIF原创"
    if as_int(source_type, -1) == 0:
        return "站内原创"
    return "转载/外部来源"


def iter_content_blocks(node: Tag) -> Iterable[Tuple[str, str]]:
    for child in node.children:
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue
        name = (child.name or "").lower()
        if name in SKIP_TAGS:
            continue
        if name == "img":
            src = (child.get("src") or "").strip()
            if src:
                yield ("img", src)
            continue
        if name in BLOCK_TAGS:
            text = normalize_text(child.get_text(" ", strip=True))
            if text:
                if name == "li":
                    text = f"• {text}"
                yield (name, text)
            for img in child.find_all("img"):
                src = (img.get("src") or "").strip()
                if src:
                    yield ("img", src)
            continue
        if name in CONTAINER_TAGS or True:
            yield from iter_content_blocks(child)


class FoodTalksExporter:
    def __init__(
        self,
        output_dir: Path,
        page_size: int = 100,
        max_pages: int = 0,
        max_details: int = 0,
        sleep_seconds: float = 0.35,
        save_every: int = 50,
        overwrite_docx: bool = False,
        refresh_list: bool = False,
        skip_docx: bool = False,
    ) -> None:
        self.output_dir = output_dir
        self.page_size = page_size
        self.max_pages = max_pages
        self.max_details = max_details
        self.sleep_seconds = sleep_seconds
        self.save_every = save_every
        self.overwrite_docx = overwrite_docx
        self.refresh_list = refresh_list
        self.skip_docx = skip_docx

        self.docs_dir = self.output_dir / "docx"
        self.docs_dir.mkdir(parents=True, exist_ok=True)

        self.jsonl_file = self.output_dir / "foodtalks_articles.jsonl"
        self.csv_file = self.output_dir / "foodtalks_articles.csv"
        self.xlsx_file = self.output_dir / "foodtalks_articles.xlsx"

        self.session = requests.Session()
        self.session.headers.update(DEFAULT_HEADERS)

        self.records: Dict[int, Dict[str, Any]] = {}

    def request_json(
        self,
        url: str,
        params: Dict[str, Any] | None = None,
        referer: str | None = None,
        retries: int = 8,
    ) -> Dict[str, Any]:
        backoff = 0.8
        last_error = ""
        for _ in range(retries):
            headers = {}
            if referer:
                headers["Referer"] = referer
            try:
                resp = self.session.get(url, params=params, headers=headers, timeout=30)
            except requests.RequestException as exc:
                last_error = str(exc)
                time.sleep(backoff + random.random() * 0.5)
                backoff = min(backoff * 1.8, 12)
                continue

            if resp.status_code == 200:
                data = resp.json()
                if data.get("code") == 0:
                    return data
                last_error = f"API code={data.get('code')} msg={data.get('msg')}"
            elif resp.status_code in {403, 429, 500, 502, 503, 504}:
                last_error = f"HTTP {resp.status_code}"
            else:
                resp.raise_for_status()

            time.sleep(backoff + random.random() * 0.5)
            backoff = min(backoff * 1.8, 20)
        raise RuntimeError(f"request failed: {url} {params or {}} {last_error}")

    def load_existing_records(self) -> None:
        if not self.jsonl_file.exists():
            return
        count = 0
        with self.jsonl_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                rec = json.loads(line)
                article_id = as_int(rec.get("id"))
                if article_id:
                    self.records[article_id] = rec
                    count += 1
        print(f"[resume] loaded existing records: {count}")

    def fetch_list_page(self, page: int) -> Dict[str, Any]:
        params = {"current": page, "size": self.page_size, "language": "ZH"}
        return self.request_json(
            LIST_API,
            params=params,
            referer=f"{SITE_BASE}/news",
        )["data"]

    def fetch_all_list(self) -> None:
        page = 1
        total_pages = None

        while True:
            data = self.fetch_list_page(page)
            records = data.get("records") or []
            total_pages = as_int(data.get("pages"), 0) or total_pages or 0
            total_count = as_int(data.get("total"), 0)

            for item in records:
                article_id = as_int(item.get("id"))
                if not article_id:
                    continue
                existing = self.records.get(article_id, {})
                merged = {
                    **existing,
                    "id": article_id,
                    "title": item.get("title") or existing.get("title") or "",
                    "publish_time": item.get("publishTime") or existing.get("publish_time") or "",
                    "author": item.get("author") or existing.get("author") or "",
                    "source_name": item.get("sourceName") or existing.get("source_name") or "",
                    "source_type": item.get("sourceType"),
                    "source_id": item.get("sourceId"),
                    "summary": item.get("summary") or existing.get("summary") or "",
                    "parent_tag_code": item.get("parentTagCode") or existing.get("parent_tag_code") or "",
                    "tag_code": item.get("tagCode") or existing.get("tag_code") or "",
                    "cover_img": item.get("coverImg") or existing.get("cover_img") or "",
                    "url": f"{SITE_BASE}/news/{article_id}",
                    "status": existing.get("status") or "pending",
                    "error": existing.get("error") or "",
                    "docx_path": existing.get("docx_path") or "",
                }
                merged["is_fbif_original"] = is_fbif_original(
                    merged.get("source_name", ""),
                    merged.get("source_type"),
                    merged.get("source_id"),
                )
                merged["source_label"] = source_label(
                    merged.get("source_name", ""),
                    merged.get("source_type"),
                    merged.get("source_id"),
                )
                self.records[article_id] = merged

            print(
                f"[list] page={page}/{total_pages or '?'} "
                f"loaded={len(records)} total_records={len(self.records)} site_total={total_count}"
            )

            if page == 1:
                self.save_records(include_excel=True)

            if self.max_pages > 0 and page >= self.max_pages:
                break
            if total_pages and page >= total_pages:
                break
            page += 1
            time.sleep(0.15)

    def fetch_detail(self, article_id: int) -> Dict[str, Any]:
        params = {
            "comment": int(time.time() * 1000),
            "language": "ZH",
        }
        referer = f"{SITE_BASE}/news/{article_id}"
        url = f"{DETAIL_API}/{article_id}"
        return self.request_json(url, params=params, referer=referer)["data"]

    def docx_output_path(self, record: Dict[str, Any]) -> Path:
        filename = f"{record['id']}_{safe_filename(record.get('title', ''))}.docx"
        return self.docs_dir / filename

    def write_docx(self, record: Dict[str, Any], detail: Dict[str, Any], target: Path) -> None:
        doc = Document()
        title = clean_docx_text(detail.get("title") or record.get("title") or "")
        if not title:
            title = f"Article {record['id']}"
        doc.add_heading(title, level=1)

        meta_lines = [
            f"文章ID: {record['id']}",
            f"发布时间: {detail.get('publishTime') or record.get('publish_time') or ''}",
            f"作者: {detail.get('author') or record.get('author') or ''}",
            f"来源: {detail.get('sourceName') or record.get('source_name') or ''}",
            f"来源分类: {record.get('source_label') or ''}",
            f"文章链接: {record.get('url') or ''}",
        ]
        for line in meta_lines:
            doc.add_paragraph(clean_docx_text(line))
        doc.add_paragraph("")

        content_html = detail.get("content") or ""
        soup = BeautifulSoup(content_html, "lxml")
        root = soup.body if soup.body else soup

        added = 0
        heading_map = {"h1": 2, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
        for tag, content in iter_content_blocks(root):
            content = clean_docx_text(content)
            if not content:
                continue
            if tag in heading_map:
                doc.add_heading(content, level=heading_map[tag])
            elif tag == "img":
                doc.add_paragraph(clean_docx_text(f"[图片] {content}"))
            elif tag == "pre":
                p = doc.add_paragraph(content)
                for run in p.runs:
                    run.font.name = "Courier New"
            else:
                doc.add_paragraph(content)
            added += 1

        if added == 0:
            fallback = clean_docx_text(normalize_text(root.get_text(" ", strip=True)))
            if fallback:
                doc.add_paragraph(fallback)

        doc.save(target)

    def process_details(self) -> None:
        pending_ids: List[int] = sorted(self.records.keys(), reverse=True)
        done = 0
        touched = 0

        for article_id in pending_ids:
            rec = self.records[article_id]
            target_path = self.docx_output_path(rec)
            rel_docx = target_path.relative_to(self.output_dir).as_posix()

            if (
                rec.get("status") == "done"
                and target_path.exists()
                and not self.overwrite_docx
            ):
                rec["status"] = "done"
                rec["docx_path"] = rel_docx
                rec["error"] = ""
                continue
            if target_path.exists() and not self.overwrite_docx:
                rec["status"] = "done"
                rec["docx_path"] = rel_docx
                rec["error"] = ""
                continue

            if self.max_details > 0 and done >= self.max_details:
                break

            try:
                detail = self.fetch_detail(article_id)
                rec["title"] = detail.get("title") or rec.get("title") or ""
                rec["publish_time"] = detail.get("publishTime") or rec.get("publish_time") or ""
                rec["author"] = detail.get("author") or rec.get("author") or ""
                rec["source_name"] = detail.get("sourceName") or rec.get("source_name") or ""
                rec["source_type"] = (
                    detail.get("sourceType")
                    if detail.get("sourceType") is not None
                    else rec.get("source_type")
                )
                rec["source_id"] = (
                    detail.get("sourceId")
                    if detail.get("sourceId") is not None
                    else rec.get("source_id")
                )
                rec["summary"] = detail.get("summary") or rec.get("summary") or ""

                rec["is_fbif_original"] = is_fbif_original(
                    rec.get("source_name", ""),
                    rec.get("source_type"),
                    rec.get("source_id"),
                )
                rec["source_label"] = source_label(
                    rec.get("source_name", ""),
                    rec.get("source_type"),
                    rec.get("source_id"),
                )

                if not self.skip_docx:
                    self.write_docx(rec, detail, target_path)
                    rec["docx_path"] = rel_docx

                rec["status"] = "done"
                rec["error"] = ""
                done += 1
                touched += 1
                print(f"[detail] ok id={article_id} done={done}")
            except Exception as exc:  # noqa: BLE001
                rec["status"] = "error"
                rec["error"] = str(exc)
                touched += 1
                print(f"[detail] fail id={article_id} error={exc}")

            if touched % self.save_every == 0:
                self.save_records(include_excel=False)
                print(f"[save] checkpoint touched={touched} done={done}")

            time.sleep(max(self.sleep_seconds, 0.0))

    def records_as_rows(self) -> List[Dict[str, Any]]:
        ordered = [self.records[k] for k in sorted(self.records.keys(), reverse=True)]
        rows: List[Dict[str, Any]] = []
        for rec in ordered:
            rows.append(
                {
                    "文章ID": rec.get("id"),
                    "标题": clean_excel_text(rec.get("title", "")),
                    "发布时间": clean_excel_text(rec.get("publish_time", "")),
                    "作者": clean_excel_text(rec.get("author", "")),
                    "来源名称": clean_excel_text(rec.get("source_name", "")),
                    "来源类型(sourceType)": rec.get("source_type"),
                    "来源ID(sourceId)": rec.get("source_id"),
                    "是否FBIF原创": "是" if rec.get("is_fbif_original") else "否",
                    "来源分类": clean_excel_text(rec.get("source_label", "")),
                    "主分类(parentTag)": clean_excel_text(rec.get("parent_tag_code", "")),
                    "子分类(tag)": clean_excel_text(rec.get("tag_code", "")),
                    "摘要": clean_excel_text(rec.get("summary", "")),
                    "文章链接": clean_excel_text(rec.get("url", "")),
                    "DOCX路径": clean_excel_text(rec.get("docx_path", "")),
                    "抓取状态": clean_excel_text(rec.get("status", "")),
                    "错误信息": clean_excel_text(rec.get("error", "")),
                }
            )
        return rows

    def save_records(self, include_excel: bool) -> None:
        ordered = [self.records[k] for k in sorted(self.records.keys(), reverse=True)]
        self.output_dir.mkdir(parents=True, exist_ok=True)

        with self.jsonl_file.open("w", encoding="utf-8") as f:
            for rec in ordered:
                f.write(json.dumps(rec, ensure_ascii=False) + "\n")

        rows = self.records_as_rows()
        df = pd.DataFrame(rows)
        df.to_csv(self.csv_file, index=False, encoding="utf-8-sig")
        if include_excel:
            df.to_excel(self.xlsx_file, index=False)

    def run(self) -> None:
        self.load_existing_records()
        if self.refresh_list or not self.records:
            self.fetch_all_list()
            self.save_records(include_excel=True)
        else:
            print("[list] using existing metadata jsonl, skip list fetch")

        if not self.skip_docx:
            self.process_details()
        self.save_records(include_excel=True)

        total = len(self.records)
        done = sum(1 for r in self.records.values() if r.get("status") == "done")
        errors = sum(1 for r in self.records.values() if r.get("status") == "error")
        print("")
        print("===== Export Finished =====")
        print(f"Total records: {total}")
        print(f"DOCX done: {done}")
        print(f"Errors: {errors}")
        print(f"JSONL: {self.jsonl_file}")
        print(f"CSV:   {self.csv_file}")
        print(f"XLSX:  {self.xlsx_file}")
        print(f"DOCX dir: {self.docs_dir}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export FoodTalks articles to DOCX + Excel")
    parser.add_argument(
        "--output-dir",
        default="output/foodtalks_export",
        help="Output directory for DOCX + metadata files",
    )
    parser.add_argument("--page-size", type=int, default=100, help="List API page size")
    parser.add_argument(
        "--max-pages",
        type=int,
        default=0,
        help="Max list pages to fetch (0 means all)",
    )
    parser.add_argument(
        "--max-details",
        type=int,
        default=0,
        help="Max article details to fetch in this run (0 means all pending)",
    )
    parser.add_argument(
        "--sleep",
        type=float,
        default=0.35,
        help="Sleep seconds between detail requests",
    )
    parser.add_argument(
        "--save-every",
        type=int,
        default=50,
        help="Save checkpoint every N touched detail records",
    )
    parser.add_argument(
        "--overwrite-docx",
        action="store_true",
        help="Overwrite existing DOCX files",
    )
    parser.add_argument(
        "--refresh-list",
        action="store_true",
        help="Ignore existing list metadata and refetch full list",
    )
    parser.add_argument(
        "--skip-docx",
        action="store_true",
        help="Only fetch/save list metadata without writing DOCX",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    exporter = FoodTalksExporter(
        output_dir=Path(args.output_dir).resolve(),
        page_size=args.page_size,
        max_pages=args.max_pages,
        max_details=args.max_details,
        sleep_seconds=args.sleep,
        save_every=args.save_every,
        overwrite_docx=args.overwrite_docx,
        refresh_list=args.refresh_list,
        skip_docx=args.skip_docx,
    )
    exporter.run()


if __name__ == "__main__":
    main()
